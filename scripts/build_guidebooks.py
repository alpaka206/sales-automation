from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "tmp" / "guide_assets"
OUT_DIR = ROOT / "output" / "guides"
TODAY = date(2026, 7, 19)

# compact_reference_guide preset. Korean glyphs use a named East Asia fallback.
BASE_FONT = "Calibri"
EAST_ASIA_FONT = "Malgun Gothic"
MONO_FONT = "Consolas"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_START_END_DXA = 120

INK = "17211F"
MUTED = "596662"
HEADING = "2E74B5"
HEADING_DARK = "1F4D78"
TABLE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
PERSO_TEAL = "367C74"       # named brand override
PERSO_SOFT = "EAF4F1"       # named brand override
CAUTION = "FFF4D6"          # named operational-warning override
RISK = "FDECEC"             # named operational-risk override
SUCCESS = "EAF6EE"          # named status override
WHITE = "FFFFFF"
LINE = "D8E0DE"

FONT_REGULAR = Path("C:/Windows/Fonts/malgun.ttf")
FONT_BOLD = Path("C:/Windows/Fonts/malgunbd.ttf")


def _rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def _set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
    italic: bool | None = None,
    font: str = BASE_FONT,
) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(
        qn("w:eastAsia"), EAST_ASIA_FONT if font == BASE_FONT else font
    )
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = _rgb(color)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_TOP_BOTTOM_DXA),
        ("bottom", CELL_TOP_BOTTOM_DXA),
        ("start", CELL_START_END_DXA),
        ("end", CELL_START_END_DXA),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _no_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_table_geometry(table, widths: Sequence[int], *, indent: int = TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        _no_row_split(row)
        for cell, width in zip(row.cells, widths, strict=True):
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(str(text))
    _set_run_font(run, size=size, bold=bold, color=color)


def _patch_numbering(doc: Document) -> None:
    """Align built-in bullet/decimal levels to compact_reference_guide tokens."""
    numbering = doc.part.numbering_part.element
    for abstract in numbering.findall(qn("w:abstractNum")):
        for lvl in abstract.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl"), "0") != "0":
                continue
            num_fmt = lvl.find(qn("w:numFmt"))
            if num_fmt is None or num_fmt.get(qn("w:val")) not in {"bullet", "decimal"}:
                continue
            p_pr = lvl.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                lvl.append(p_pr)
            ind = p_pr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                p_pr.append(ind)
            ind.set(qn("w:left"), "540")
            ind.set(qn("w:hanging"), "270")
            tabs = p_pr.find(qn("w:tabs"))
            if tabs is None:
                tabs = OxmlElement("w:tabs")
                p_pr.append(tabs)
            for old in list(tabs):
                tabs.remove(old)
            tab = OxmlElement("w:tab")
            tab.set(qn("w:val"), "num")
            tab.set(qn("w:pos"), "540")
            tabs.append(tab)


def _add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    _set_run_font(run, size=8.5, color=MUTED)


def _configure_document(title: str, short_label: str) -> Document:
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, HEADING, 18, 10),
        ("Heading 2", 13, HEADING, 14, 7),
        ("Heading 3", 12, HEADING_DARK, 10, 5),
    ):
        style = styles[name]
        style.font.name = BASE_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = _rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BASE_FONT
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn("w:ascii"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BASE_FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    caption = styles["Caption"]
    caption.font.name = BASE_FONT
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = _rgb(MUTED)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _patch_numbering(doc)

    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        left = p.add_run("PERSO INBOUND  |  ")
        _set_run_font(left, size=8.5, bold=True, color=PERSO_TEAL)
        right = p.add_run(short_label)
        _set_run_font(right, size=8.5, color=MUTED)

    for footer in (section.footer, section.even_page_footer):
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        fp.paragraph_format.space_before = Pt(0)
        _add_page_field(fp)

    # Materialize explicit empty first-page parts for the cover. Word's PDF
    # exporter otherwise applies the title-page fallback inconsistently when
    # odd/even running headers are both present.
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    doc.core_properties.title = title
    doc.core_properties.subject = "PERSO inbound sales automation guide"
    doc.core_properties.author = "PERSO"
    doc.core_properties.keywords = "inbound, HubSpot, Google Sheets, automation, operations"
    doc.core_properties.comments = "Generated from the verified local implementation."
    return doc


def add_para(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    color: str = INK,
    size: float = 11,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float = 6,
    italic: bool = False,
) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    _set_run_font(r, size=size, bold=bold, color=color, italic=italic)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    if level == 1 and getattr(doc, "_perso_next_page_break", False):
        p.paragraph_format.page_break_before = True
        setattr(doc, "_perso_next_page_break", False)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    _set_run_font(
        r,
        size={1: 16, 2: 13, 3: 12}[level],
        bold=True,
        color={1: HEADING, 2: HEADING, 3: HEADING_DARK}[level],
    )


def add_bullets(doc: Document, items: Iterable[str], *, size: float = 11) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        _set_run_font(r, size=size, color=INK)


def add_steps(doc: Document, items: Iterable[str], *, size: float = 11) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        number = p.add_run(f"{index}. ")
        _set_run_font(number, size=size, bold=True, color=HEADING_DARK)
        r = p.add_run(item)
        _set_run_font(r, size=size, color=INK)


def add_callout(
    doc: Document,
    title: str,
    body: str,
    *,
    fill: str = CALLOUT_FILL,
    accent: str = PERSO_TEAL,
) -> None:
    table = doc.add_table(rows=1, cols=1)
    _set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    _shade_cell(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "18")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    _set_run_font(r, size=10.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    _set_run_font(r2, size=10, color=INK)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    widths: Sequence[int],
    *,
    font_size: float = 9.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    _set_table_geometry(table, widths)
    _repeat_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        _shade_cell(cell, TABLE_FILL)
        _set_cell_text(cell, header, bold=True, color=HEADING_DARK, size=9.3)
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            _set_cell_text(cells[idx], str(value), color=INK, size=font_size)
        _set_table_geometry(table, widths)
    add_para(doc, "", after=2)


def add_code_block(doc: Document, lines: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    _set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    _shade_cell(cell, "F2F4F7")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(lines.splitlines()):
        if idx:
            p.add_run().add_break()
        r = p.add_run(line)
        _set_run_font(r, size=8.5, color=INK, font=MONO_FONT)
    add_para(doc, "", after=2)


def add_image(doc: Document, path: Path, caption: str, alt: str, *, width: float = 6.3) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", caption)
    cp = doc.add_paragraph(style="Caption")
    cr = cp.add_run(caption)
    _set_run_font(cr, size=9, color=MUTED)


def add_page_break(doc: Document) -> None:
    # Defer the break to the next top-level heading. A standalone page-break
    # paragraph can be pushed onto an otherwise blank page when the prior page
    # is exactly full in Word.
    setattr(doc, "_perso_next_page_break", True)


def add_cover(doc: Document, *, kicker: str, title: str, subtitle: str, audience: str) -> None:
    logo = ROOT / "src" / "api" / "web" / "static" / "logo.png"
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(72)
    p_logo.paragraph_format.space_after = Pt(22)
    if logo.exists():
        shape = p_logo.add_run().add_picture(str(logo), width=Inches(0.72))
        shape._inline.docPr.set("descr", "PERSO 로고")
    add_para(
        doc,
        kicker.upper(),
        size=10,
        bold=True,
        color=PERSO_TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=16,
    )
    add_para(
        doc,
        title,
        size=29,
        bold=True,
        color=INK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=10,
    )
    add_para(
        doc,
        subtitle,
        size=14,
        color=HEADING_DARK,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=36,
    )
    add_para(
        doc,
        audience,
        size=10.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=48,
    )
    add_para(
        doc,
        f"검증 기준일  {TODAY.isoformat()}  |  로컬 구현 기준",
        size=9.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=4,
    )
    add_para(
        doc,
        "PERSO Inbound Operations",
        size=10,
        bold=True,
        color=PERSO_TEAL,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def add_status_legend(doc: Document) -> None:
    add_table(
        doc,
        ["표시", "뜻", "운영자가 할 일"],
        [
            ("구현됨", "현재 코드와 화면에서 동작", "계정·환경값만 정상인지 확인"),
            ("연결 필요", "기능은 있으나 실제 외부 계정 권한이 필요", "관리자 또는 개발자가 1회 연결"),
            ("부분 구현", "기록·관리까지만 가능", "외부 서비스 동작은 사람이 수행"),
            ("미구현", "화면/자동화가 아직 없음", "수동 절차를 쓰고 백로그로 관리"),
        ],
        [1440, 3600, 4320],
    )


# ----------------------------- diagrams -----------------------------


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_BOLD if bold else FONT_REGULAR), size=size)


def _wrap(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for token in text.split(" "):
        candidate = token if not current else f"{current} {token}"
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = token
    if current:
        lines.append(current)
    return lines


def _box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    title: str,
    body: str,
    *,
    fill: str = "FFFFFF",
    outline: str = LINE,
    accent: str = PERSO_TEAL,
) -> None:
    draw.rounded_rectangle(xy, radius=22, fill=f"#{fill}", outline=f"#{outline}", width=2)
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle((x1 + 18, y1 + 18, x1 + 28, y2 - 18), radius=5, fill=f"#{accent}")
    draw.text((x1 + 48, y1 + 25), title, font=_font(28, True), fill=f"#{INK}")
    body_lines = _wrap(draw, body, _font(20), x2 - x1 - 75)
    y = y1 + 72
    for line in body_lines[:3]:
        draw.text((x1 + 48, y), line, font=_font(20), fill=f"#{MUTED}")
        y += 30


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line((start, end), fill=f"#{PERSO_TEAL}", width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        sign = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * sign, ey - 10), (ex - 18 * sign, ey + 10)]
    else:
        sign = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 10, ey - 18 * sign), (ex + 10, ey - 18 * sign)]
    draw.polygon(points, fill=f"#{PERSO_TEAL}")


def build_diagrams() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)

    img = Image.new("RGB", (1600, 900), "#F7F9F8")
    d = ImageDraw.Draw(img)
    d.text((70, 50), "문의가 들어와 답변이 나가기까지", font=_font(42, True), fill=f"#{INK}")
    d.text((70, 105), "접수확인과 상세답변은 서로 다른 흐름입니다.", font=_font(23), fill=f"#{MUTED}")
    boxes = [
        ((70, 190, 420, 360), "1. New 문의", "HubSpot Webhook, 누락 시 10분 폴링"),
        ((625, 190, 975, 360), "2. 즉시 접수확인", "첫 문의에만, 사람 승인 없이 고객 언어로 발송"),
        ((1180, 190, 1530, 360), "3. AI + 내부 규정", "문의 분류, 관련 문서 선택, 한국어 검토 초안"),
        ((1180, 590, 1530, 760), "4. 사람 검토", "수정, 번역, 서명, 미리보기 후 발송 승인"),
        ((625, 590, 975, 760), "5. 실제 발송", "SMTP 이메일로 발송"),
        ((70, 590, 420, 760), "6. 상태 동기화", "HubSpot, 사이트 파이프라인, Inbound DB 갱신"),
    ]
    for xy, title, body in boxes:
        _box(d, xy, title, body)
    _arrow(d, (420, 275), (625, 275))
    _arrow(d, (975, 275), (1180, 275))
    _arrow(d, (1355, 360), (1355, 590))
    _arrow(d, (1180, 675), (975, 675))
    _arrow(d, (625, 675), (420, 675))
    img.save(ASSET_DIR / "nondev_flow.png")

    img = Image.new("RGB", (1600, 860), "#F7F9F8")
    d = ImageDraw.Draw(img)
    d.text((70, 50), "왼쪽 메뉴는 세 가지 일로 나뉩니다", font=_font(42, True), fill=f"#{INK}")
    cols = [
        (70, "인바운드 답장", ["문의 대시보드", "답변 검토", "고객·회사 히스토리", "이메일 규칙·서명", "정책·지식 문서"]),
        (560, "인사이트", ["문의·국가 추이", "업데이트 필요 고객", "답장 누락", "갱신 임박", "장애 복구"]),
        (1050, "파이프라인 연동관리", ["문의 파이프라인", "HubSpot 단계", "Google Sheets 연결", "견적·결제 준비", "수주 DB 반영"]),
    ]
    fills = ["EAF4F1", "EEF2F8", "F4F1EA"]
    for (x, title, items), fill in zip(cols, fills, strict=True):
        d.rounded_rectangle((x, 165, x + 430, 785), radius=28, fill=f"#{fill}", outline=f"#{LINE}", width=2)
        d.text((x + 35, 205), title, font=_font(30, True), fill=f"#{INK}")
        y = 285
        for index, item in enumerate(items, start=1):
            d.ellipse((x + 35, y + 3, x + 63, y + 31), fill=f"#{PERSO_TEAL}")
            d.text((x + 43, y + 2), str(index), font=_font(16, True), fill="#FFFFFF")
            d.text((x + 82, y), item, font=_font(23), fill=f"#{INK}")
            y += 88
    img.save(ASSET_DIR / "menu_map.png")

    img = Image.new("RGB", (1600, 900), "#F7F9F8")
    d = ImageDraw.Draw(img)
    d.text((70, 50), "시스템 아키텍처 한 장 요약", font=_font(42, True), fill=f"#{INK}")
    d.text((70, 105), "외부 이벤트는 영속 큐를 거쳐 처리되며, 고객 발송과 외부 동기화는 분리됩니다.", font=_font(22), fill=f"#{MUTED}")
    _box(d, (70, 195, 430, 365), "HubSpot", "Webhook 우선, 10분 poller 보완", fill="FFFFFF")
    _box(d, (620, 195, 980, 365), "FastAPI + InboundJob", "서명 검증, 중복제거, DB 영속 큐", fill="EAF4F1")
    _box(d, (1170, 195, 1530, 365), "Inbound worker", "owner lease+heartbeat, AI 분석, 초안", fill="FFFFFF")
    _box(d, (1170, 580, 1530, 760), "Send worker", "atomic claim, DB quota, SMTP, retry", fill="FFFFFF")
    _box(d, (620, 580, 980, 760), "Post-send sync", "HubSpot 활동/단계, Sheets 단계", fill="EAF4F1")
    _box(d, (70, 580, 430, 760), "PostgreSQL", "시스템 원장: 문의, 메시지, 계약, 감사", fill="FFFFFF")
    _arrow(d, (430, 280), (620, 280))
    _arrow(d, (980, 280), (1170, 280))
    _arrow(d, (1350, 365), (1350, 580))
    _arrow(d, (1170, 670), (980, 670))
    _arrow(d, (620, 670), (430, 670))
    d.text((515, 430), "Gemini Vertex AI  |  정책·지식 문서  |  승인 UI·장애 복구", font=_font(24, True), fill=f"#{HEADING_DARK}")
    img.save(ASSET_DIR / "dev_architecture.png")

    img = Image.new("RGB", (1600, 620), "#F7F9F8")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "문의 파이프라인 단계", font=_font(42, True), fill=f"#{INK}")
    stages = ["New", "Meeting link sent", "Negotiation", "Contracted", "Onboarding", "Active"]
    x = 55
    for idx, stage in enumerate(stages):
        w = 220
        fill = "EAF4F1" if idx in {0, 1} else "FFFFFF"
        d.rounded_rectangle((x, 210, x + w, 350), radius=22, fill=f"#{fill}", outline=f"#{LINE}", width=2)
        lines = stage.split(" ")
        y = 245 if len(lines) == 1 else 230
        for line in lines:
            bbox = d.textbbox((0, 0), line, font=_font(22, True))
            d.text((x + (w - (bbox[2] - bbox[0])) / 2, y), line, font=_font(22, True), fill=f"#{INK}")
            y += 34
        if idx < len(stages) - 1:
            _arrow(d, (x + w, 280), (x + w + 38, 280))
        x += 258
    d.rounded_rectangle((1180, 455, 1530, 555), radius=20, fill=f"#{RISK}", outline="#E6B6B6", width=2)
    d.text((1245, 488), "Closed Lost", font=_font(25, True), fill=f"#{INK}")
    d.text((70, 455), "상세답변 발송 성공 시 자동으로 두 번째 단계로 이동합니다. 이후 단계는 카드 이동 또는 계약 상태 저장으로 관리합니다.", font=_font(22), fill=f"#{MUTED}")
    img.save(ASSET_DIR / "status_flow.png")

    img = Image.new("RGB", (1600, 720), "#F7F9F8")
    d = ImageDraw.Draw(img)
    d.text((70, 45), "장애 복구는 재발송보다 상태 확인이 먼저입니다", font=_font(40, True), fill=f"#{INK}")
    d.text((70, 100), "사이트의 장애 복구 화면은 안전하게 다시 할 수 있는 작업과 사람 확인이 필요한 작업을 분리합니다.", font=_font(21), fill=f"#{MUTED}")
    recovery_boxes = [
        ((70, 190, 410, 430), "Dead 문의 작업", "원인 수정\n→ 문의 작업 재시도"),
        ((445, 190, 785, 430), "확정 발송 실패", "주소·SMTP 수정\n→ 발송 재시도"),
        ((820, 190, 1160, 430), "발송 확인 필요", "보낸메일함 확인\n→ 보냄/안 보냄 확정"),
        ((1195, 190, 1535, 430), "외부 동기화 실패", "연결 복구\n→ sync만 재시도"),
    ]
    for xy, title, body in recovery_boxes:
        _box(d, xy, title, body, fill="FFFFFF")
    d.rounded_rectangle((250, 520, 1350, 640), radius=24, fill=f"#{RISK}", outline="#E6B6B6", width=2)
    d.text((325, 550), "delivery_unknown은 확인 없이 재발송하지 않습니다. 중복 메일을 막는 가장 중요한 규칙입니다.", font=_font(24, True), fill=f"#{INK}")
    img.save(ASSET_DIR / "recovery_flow.png")


def save_document(doc: Document, filename: str) -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUT_DIR / filename
    doc.save(path)
    return path


def build_developer_guide() -> Path:
    doc = _configure_document(
        "PERSO 인바운드 자동화 - 개발자 설정 및 아키텍처 가이드",
        "개발자 설정·아키텍처",
    )
    add_cover(
        doc,
        kicker="Engineering Handbook",
        title="개발자 설정 및\n아키텍처 가이드",
        subtitle="설치, 연동, 보안, 배포, 장애 대응까지",
        audience="대상: 백엔드·플랫폼·DevOps·보안·운영 관리자",
    )

    add_page_break(doc)
    add_heading(doc, "1. 문서 목적과 현재 상태", 1)
    add_para(
        doc,
        "이 문서는 현재 저장소의 실제 코드와 테스트를 기준으로, 새로운 개발자가 시스템을 설치하고 외부 계정을 연결하며 실서비스로 전환할 수 있도록 작성되었다. 운영 화면의 설명은 별도 비개발자 가이드에 있다.",
    )
    add_status_legend(doc)
    add_heading(doc, "현재 개발 환경 스냅샷", 2)
    add_table(
        doc,
        ["영역", "현재 상태", "실서비스 전 조치"],
        [
            ("코드·테스트", "Ruff 통과, 전체 521개 테스트 통과", "CI의 Python 3.11/3.12·PostgreSQL·Docker gate 유지"),
            ("HubSpot", "개발 계정 New=1, 발송 후=2", "실제 pipeline stage 내부 ID로 교체"),
            ("Google Sheets", "서비스 계정 자격증명 감지, 편집 권한 불가", "사용자 OAuth 웹 클라이언트 생성·연결"),
            ("SMTP", "자격증명 있음, SMTP가 실제 전달 수단", "SEND_OVERRIDE_EMAIL로 1건 검증 후 해제"),
            ("Slack", "토큰은 있을 수 있으나 전체 스위치와 채널 모드 OFF", "필요할 때 reply-ready 알림만 활성"),
            ("UI 인증", "로컬 basic 모드", "운영은 회사 Google OAuth와 allowlist 권장"),
            ("배포", "단일 프로세스·worker heartbeat·시작 안전검사", "Managed PostgreSQL, 공개 HTTPS, WEB_CONCURRENCY=1"),
        ],
        [1800, 3600, 3960],
        font_size=8.8,
    )
    add_page_break(doc)
    add_heading(doc, "2. 아키텍처 한 장 요약", 1)
    add_image(
        doc,
        ASSET_DIR / "dev_architecture.png",
        "그림 1. Webhook 수신부터 발송 후 동기화까지의 시스템 경계",
        "HubSpot, FastAPI 영속 큐, inbound worker, send worker, post-send sync, PostgreSQL로 이어지는 시스템 아키텍처",
    )
    add_table(
        doc,
        ["원칙", "구현 의미"],
        [
            ("Webhook은 빨리 응답", "서명 검증과 DB enqueue까지만 하고 외부 API/AI 호출은 worker로 넘긴다."),
            ("DB가 시스템 원장", "Google Sheets는 운영 미러이며 백업이나 유일한 정합성 원장이 아니다."),
            ("고객 발송과 CRM 동기화 분리", "SMTP 성공 후 HubSpot/Sheets가 실패해도 메일을 실패로 되돌리거나 재발송하지 않는다."),
            ("안전 기본값", "상세 답변은 기본적으로 사람 승인, Slack은 기본 비활성이다."),
        ],
        [2160, 7200],
    )

    add_page_break(doc)
    add_heading(doc, "3. End-to-end 처리 순서", 1)
    add_steps(
        doc,
        [
            "HubSpot ticket.creation 또는 hs_pipeline_stage가 New가 되는 propertyChange를 수신한다.",
            "v3 HMAC, timestamp, 1MiB body, 100-event 제한을 검증한 뒤 InboundJob에 idempotent enqueue한다.",
            "Inbound worker가 티켓·연락처·문의 본문을 조회하고 New 단계인지 다시 확인한다.",
            "문의·Conversation·drafting placeholder를 먼저 DB에 저장해 UI에 즉시 보이게 한다.",
            "첫 문의이면 auto_ack를 일반 send queue로 승인 상태에 넣고 즉시 발송을 시도한다.",
            "Inbound DB Client ID를 예약하고 시트 미러를 시도한다. 실패해도 poller가 backfill한다.",
            "Gemini가 언어, 분류, 점수, 관련 활성 정책 문서를 선택하고 한국어 검토 초안을 작성한다.",
            "기본값은 pending_approval. 이 시점에만 Slack 알림을 보낼 수 있다.",
            "운영자가 수정·번역·서명을 선택하고 승인하면 send worker가 atomic claim 후 SMTP로 발송한다.",
            "상세 메일 성공을 먼저 커밋하고, HubSpot 활동·단계와 Inbound DB 단계는 별도 retry로 갱신한다.",
        ],
    )
    add_callout(
        doc,
        "중요한 분리",
        "auto_ack는 접수 확인일 뿐 파이프라인을 이동하지 않는다. Meeting link sent는 상세 답변 발송 성공 시점의 업무 상태이며, 실제 메일에 미팅 링크가 포함됐는지를 검사하는 상태는 아니다.",
    )

    add_page_break(doc)
    add_heading(doc, "4. 저장소 구조와 책임", 1)
    add_table(
        doc,
        ["경로", "책임"],
        [
            ("src/api/main.py", "FastAPI 진입점, 인증 middleware, health, background task 수명주기"),
            ("src/api/webhook.py", "HubSpot 서명 검증, payload 제한, durable enqueue"),
            ("src/agents/inbound_worker.py", "InboundJob owner lease·heartbeat, retry, dead 처리"),
            ("src/agents/inbound.py", "문의 저장, auto-ack, Sheets mirror, AI 분류·초안"),
            ("src/agents/send_worker.py", "승인 메시지 claim, rate limit, SMTP, delivery_unknown, post-sync retry"),
            ("src/agents/sheet_sync.py", "Inbound DB import/backfill, 수주 DB sync, durable 수동 sync 요청"),
            ("src/integrations", "HubSpot, Google Sheets/OAuth, SMTP, Slack, HTML 이메일"),
            ("src/api/web", "운영자 UI routes, templates, auth, static design system"),
            ("src/db/models.py", "연락처·대화·메시지·계약·문서·자격증명 데이터 모델"),
            ("src/db/migrations", "기존 데이터를 보존하는 additive schema migration"),
            ("company_rules / knowledge_base", "AI 답변 원칙과 선택 가능한 제품·정책 문서 seed"),
            ("tests", "외부 네트워크를 mock한 hermetic unit/integration/e2e 테스트"),
        ],
        [2900, 6460],
        font_size=8.7,
    )
    add_heading(doc, "핵심 진입점", 2)
    add_code_block(
        doc,
        "Web:     uvicorn src.api.main:app --host 127.0.0.1 --port 8000\n"
        "Migrate: python -m src.db.migrate\n"
        "Doctor:  python -m src.cli doctor\n"
        "Tests:   python -m pytest -q\n"
        "Lint:    python -m ruff check .",
    )

    add_page_break(doc)
    add_heading(doc, "5. 로컬 설치와 실행", 1)
    add_heading(doc, "Windows 권장 경로", 2)
    add_steps(
        doc,
        [
            "Python 3.11 이상을 설치한다.",
            "저장소 루트에서 scripts\\setup.bat을 실행해 .venv와 개발 의존성을 만든다.",
            ".env.example을 기준으로 .env를 채운다. 기존 .env를 덮어쓰지 않는다.",
            "python scripts/init_db.py 또는 python -m src.db.migrate를 실행한다.",
            "python -m src.cli doctor로 필수 설정과 연결 가능성을 확인한다.",
            "scripts\\run.bat을 실행하고 http://127.0.0.1:8000/healthz와 /를 확인한다.",
        ],
    )
    add_code_block(
        doc,
        "scripts\\setup.bat\n"
        ".\\.venv\\Scripts\\python.exe -m src.db.migrate\n"
        ".\\.venv\\Scripts\\python.exe -m src.cli doctor\n"
        "scripts\\run.bat",
    )
    add_heading(doc, "안전한 라이브 검증", 2)
    add_bullets(
        doc,
        [
            "SEND_OVERRIDE_EMAIL에 내부 테스트 주소를 설정하면 원본 DB 수신자를 바꾸지 않고 테스트 주소로만 보내며 상태를 test_sent로 기록한다.",
            "test_sent는 HubSpot 활동·문의 단계·Sheets를 전혀 변경하지 않는다.",
            "실서비스 전 반드시 SEND_OVERRIDE_EMAIL을 비운다.",
            "개발 계정 단계 ID 1→2를 실제 계정에 그대로 복사하지 않는다. scripts/list_ticket_stages.py로 내부 ID를 확인한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "6. 환경변수 설정 지도", 1)
    add_table(
        doc,
        ["분류", "필수/대표 키", "검증"],
        [
            ("앱·DB", "DATABASE_URL, INTERNAL_API_TOKEN, APP_HOST/PORT, TIMEZONE", "healthz, doctor"),
            ("Gemini", "GOOGLE_CREDENTIALS_JSON, GOOGLE_CLOUD_PROJECT/LOCATION, model 2종", "internal healthcheck"),
            ("HubSpot", "PRIVATE_APP_TOKEN, WEBHOOK_SECRET, stage ID 7종", "list_ticket_stages, test ticket"),
            ("SMTP", "HOST/PORT/USERNAME/PASSWORD/FROM", "override 주소 1건"),
            ("Inbound", "POLL/WORKER/AUTO_ACK, 600초, 24h lookback", "logs, InboundJob"),
            ("발송", "SEND_WORKER, RATE, DAILY_LIMIT, JITTER, AUTO_SEND_THRESHOLD", "approved→sent"),
            ("Google Sheets", "OAuth client ID/secret, SESSION_SECRET, GOOGLE_TOKEN_ENCRYPTION_KEY", "pipeline 연결 패널"),
            ("Slack", "SLACK_ENABLED, APPROVAL_CHANNEL, token, channel ID", "pending approval 1건"),
            ("UI 인증", "AUTH_MODE, Google OAuth, allowed domain/list, SESSION_SECRET", "viewer/operator/admin 로그인"),
            ("프로세스", "WEB_CONCURRENCY=1, worker enable 3종", "startup guard, heartbeat health"),
        ],
        [1680, 5000, 2680],
        font_size=8.4,
    )
    add_callout(
        doc,
        "자동발송 안전값",
        "AUTO_SEND_THRESHOLD=1.01은 모든 상세 답변에 사람 승인을 요구한다. 0.0~1.0은 score/100이 임계치 이상일 때 자동승인을 켜므로 파일럿 검증 전에는 사용하지 않는다.",
        fill=CAUTION,
        accent="9A6A00",
    )

    add_page_break(doc)
    add_heading(doc, "7. 데이터 모델과 마이그레이션", 1)
    add_table(
        doc,
        ["모델", "역할", "중요 키/규칙"],
        [
            ("Contact", "사람·회사 식별", "normalized_email unique, domain, HubSpot ID, sheet client ID"),
            ("Conversation", "한 문의/티켓의 thread", "HubSpot ticket ID unique, stage, language, summary, sheet key"),
            ("Message", "inbound/auto_ack/draft/reply", "status, target language, send lease, post-sync, Slack flags"),
            ("InboundJob", "Webhook/poller durable work", "event_key unique, owner lease·heartbeat, attempts, available_at"),
            ("CustomerProfile", "운영자 관리 상태", "pipeline, temperature, MQL/PQL, next action, plan"),
            ("CustomerInteraction", "다채널 통합 이력", "meeting/Kakao/phone/manual, artifact URL"),
            ("ContractRecord", "계약·결제·갱신", "문의 conversation snapshot, Decimal 금액, order sheet sync"),
            ("Knowledge/Email revisions", "정책·템플릿 버전", "수정 전 snapshot을 append-only 보관"),
            ("IntegrationCredential", "위임 OAuth", "provider별 1행, refresh token 암호화"),
        ],
        [1800, 3240, 4320],
        font_size=8.4,
    )
    add_heading(doc, "운영 DB 원칙", 2)
    add_bullets(
        doc,
        [
            "로컬은 SQLite WAL을 허용하지만 운영은 managed PostgreSQL을 사용한다.",
            "Google Sheets는 미러이며 복구 원장이 아니다.",
            "마이그레이션 전 snapshot을 만들고, scripts/init_db.py 또는 src.db.migrate를 배포 단계에서 1회 실행한다.",
            "0032 unique index는 오래된 운영 DB에 중복 email/ticket이 있으면 실패할 수 있으므로 사전 중복 감사 후 실행한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "8. Queue, idempotency, retry", 1)
    add_table(
        doc,
        ["구간", "보장", "실패 시"],
        [
            ("Webhook→InboundJob", "event_key unique, 짧은 transaction", "동일 event는 duplicate, HubSpot retry 안전"),
            ("Inbound worker", "owner token claim, 30분 lease+10분 heartbeat, 최대 8회 retry", "소진 시 dead, 장애 복구에서 재처리"),
            ("auto_ack", "일반 Message claim 재사용, thread당 1건", "transient 최대 5회, 상세 답변은 계속"),
            ("상세 발송", "approved→sending:<worker> atomic claim", "SMTP transient 3회, permanent는 send_failed"),
            ("stale sending", "자동 replay 금지", "15분 후 delivery_unknown 격리"),
            ("post-send sync", "메일 sent 커밋 후 외부 상태만 재시도", "최대 8회, 고객 메일 재전송 없음"),
            ("Sheets", "Client ID 선예약, key upsert, 수동 요청을 Event에 영속 저장", "poller가 import/backfill 처리, 실패 상태 표시"),
        ],
        [1800, 4140, 3420],
        font_size=8.6,
    )
    add_callout(
        doc,
        "SMTP exactly-once 한계",
        "DB claim은 원자적이지만 SMTP 서버가 메일을 수락한 직후 연결이 끊기면 성공 여부를 확정할 수 없다. delivery_unknown은 보낸메일함과 제공자 로그의 Message-ID를 확인한 뒤에만 수동 판단하며 즉시 재발송하지 않는다.",
        fill=RISK,
        accent="9B1C1C",
    )

    add_page_break(doc)
    add_heading(doc, "9. HubSpot 연결", 1)
    add_steps(
        doc,
        [
            "Private App에 연락처·티켓·활동 읽기/쓰기 권한을 최소 범위로 부여한다.",
            "Webhook 앱에 ticket.creation과 ticket.propertyChange(hs_pipeline_stage)를 등록한다.",
            "Callback은 https://<PUBLIC_BASE_URL>/webhooks/hubspot 로 설정한다.",
            "HUBSPOT_WEBHOOK_SECRET과 signature required=true를 운영에서 고정한다.",
            "실제 pipeline의 New, Meeting link sent, Negotiation, Contracted, Onboarding, Active, Closed Lost 내부 ID를 환경값에 매핑한다.",
            "문의 본문·연락처 association이 있는 내부 테스트 티켓을 New로 이동해 end-to-end를 검증한다.",
        ],
    )
    add_heading(doc, "수신 안정성", 2)
    add_bullets(
        doc,
        [
            "Webhook은 1차 경로다. poller는 hs_lastmodifieddate 기준 15분 overlap으로 기본 10분마다 누락을 보완한다.",
            "다른 단계에서 생성됐다가 New로 이동한 티켓도 propertyChange 또는 poller가 잡는다.",
            "TLS proxy가 Host/X-Forwarded-Host/Proto를 신뢰 가능한 값으로 재설정하도록 하고 앱 직접 접근을 차단한다.",
            "WAF/reverse proxy에서도 body size와 rate limit을 추가하고 webhook 401/413/5xx를 경보한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "10. Google Sheets 연결과 정합성", 1)
    add_heading(doc, "권장: 관리자 사용자 OAuth", 2)
    add_steps(
        doc,
        [
            "Google Cloud에서 OAuth 2.0 Web application client를 만든다.",
            "로컬 callback http://127.0.0.1:8000/integrations/google-sheets/callback 과 운영 callback을 승인 URI에 등록한다.",
            "Client ID/Secret을 GOOGLE_SHEETS_OAUTH_CLIENT_ID/SECRET에 저장하고 서버를 재시작한다.",
            "SESSION_SECRET과 별도의 GOOGLE_TOKEN_ENCRYPTION_KEY를 secret manager에 생성한다.",
            "파이프라인 화면에서 시트 편집 권한이 있는 전용 업무 계정으로 연결한다.",
            "시트 읽기·동기화를 누르면 요청이 DB에 먼저 저장된다. 처리 중·완료·실패 상태를 화면에서 확인한다.",
        ],
    )
    add_table(
        doc,
        ["탭", "방향", "보존 규칙"],
        [
            ("Inbound DB", "신규 문의 append/upsert, 단계 일부 update, 수동 import", "기존 헤더·수식·서식 변경 금지"),
            ("Inbound 퀄리티 분석", "앱에서 쓰지 않음", "수식·차트 보호. 앱 인사이트는 로컬 DB 집계"),
            ("수주 DB", "contracted/active 계약 upsert", "Client ID+수주일 key, 기존 헤더 매핑"),
        ],
        [2100, 3600, 3660],
    )
    add_callout(
        doc,
        "단일 writer 제약",
        "Google Sheets의 조회 후 append는 외부 transaction이 아니고 _APPEND_LOCK은 한 프로세스에서만 유효하다. 현재 배포는 단일 Sheets writer를 유지한다. 수평 확장 전 advisory lock 또는 전용 sync queue를 도입한다.",
        fill=CAUTION,
        accent="9A6A00",
    )

    add_page_break(doc)
    add_heading(doc, "11. 이메일과 Slack", 1)
    add_table(
        doc,
        ["채널", "현재 동작", "운영 주의"],
        [
            ("SMTP", "실제 고객 이메일 전달, TLS, stable Message-ID, HTML 서명", "rate/day cap, delivery_unknown runbook"),
            ("HubSpot Email", "발송 성공 후 CRM 활동 기록", "실제 전달 수단이 아님"),
            ("Slack", "pending_approval 초안당 1회, 실패만 최대 5회 재시도", "비공개 승인 채널, PII 보존정책"),
        ],
        [1560, 4680, 3120],
    )
    add_heading(doc, "템플릿과 언어 안전장치", 2)
    add_bullets(
        doc,
        [
            "auto_ack는 {name}을 치환하고 고객 문의 언어로 번역한다. 번역 실패 시 한국어를 유지한다.",
            "상세 초안은 한국어 검토용이며 target_language를 별도 보존한다. 발송 직전 언어 guard가 다시 확인한다.",
            "첫 상세 회신의 가격 문장은 draft와 send 양쪽 guard에서 제한한다.",
            "HTML 서명 이미지에는 공개 HTTPS URL만 사용한다. 로컬 경로와 인증 필요 URL은 금지한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "12. 인증, 권한, 개인정보", 1)
    add_table(
        doc,
        ["영역", "현재 구현", "실서비스 강화"],
        [
            ("API", "X-Internal-Token, per-message approval HMAC", "secret manager, rotation"),
            ("웹 UI", "localhost/basic 또는 회사 Google OAuth", "운영은 Google OAuth+allowlist"),
            ("역할", "viewer=조회, operator=업무 변경, admin=설정·연동·정책", "최소권한 정기 검토"),
            ("세션", "7일 signed cookie, 매 요청 사용자 DB 재검증", "SESSION_SECRET rotation runbook"),
            ("CSRF", "SameSite=Lax + Origin/Referer·Sec-Fetch 검사", "reverse proxy에서도 허용 origin 고정"),
            ("OAuth token", "전용 키 Fernet 암호화 DB 저장", "GOOGLE_TOKEN_ENCRYPTION_KEY 별도 backup"),
            ("로그/Slack", "이메일·전화·대표 secret 패턴 자동 마스킹", "접근 최소화, 보존기간"),
            ("API docs", "localhost에서만 /docs·openapi 노출", "운영에서 직접 앱 포트 비공개"),
        ],
        [1560, 3900, 3900],
        font_size=8.5,
    )
    add_callout(
        doc,
        "권한 경계",
        "operator는 답변·고객·계약·파이프라인·장애 복구를 처리한다. 정책·이메일 규칙·Google 연결·사용자·로그 변경은 admin만 가능하며 viewer의 변경 요청은 서버에서 거부된다.",
        fill=SUCCESS,
        accent=PERSO_TEAL,
    )

    add_page_break(doc)
    add_heading(doc, "13. 배포와 확장", 1)
    add_heading(doc, "단일 인스턴스 기준", 2)
    add_steps(
        doc,
        [
            "Managed PostgreSQL을 만들고 DATABASE_URL과 자동 backup/PITR을 활성화한다.",
            "pyproject extras 설치 후 migration을 실행한다.",
            "PUBLIC_BASE_URL, UI 인증, 모든 secret을 배포 환경에 넣는다.",
            "uvicorn src.api.main:app를 단일 process로 시작한다.",
            "healthz, internal healthcheck, logs, operations, worker 처리와 외부 연동을 확인한다.",
        ],
    )
    add_callout(
        doc,
        "수평 확장 전 필수",
        "현재 각 web process가 inbound/poller/send worker를 함께 시작하며 startup guard가 worker 사용 중 WEB_CONCURRENCY>1을 거부한다. web과 worker를 분리하거나 leader election을 구현하기 전에는 여러 인스턴스를 띄우지 않는다.",
        fill=CAUTION,
        accent="9A6A00",
    )
    add_bullets(
        doc,
        [
            "Docker 이미지는 non-root 사용자·healthcheck·PostgreSQL extra를 포함하고 migration은 배포 pipeline에서 별도 실행한다.",
            "CI는 Python 3.11/3.12 테스트, pip check/audit, PostgreSQL 16 migration smoke, Docker build를 gate로 둔다.",
            "발송량은 DB의 sent_at 기준으로 재계산하지만 다중 인스턴스 전역 동시성은 지원하지 않으므로 WEB_CONCURRENCY=1을 유지한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "14. 모니터링, 백업, 테스트", 1)
    add_heading(doc, "권장 경보", 2)
    add_table(
        doc,
        ["조건", "우선순위", "행동"],
        [
            ("delivery_unknown > 0", "즉시", "제공자 Sent/Message-ID 확인, 자동 재발송 금지"),
            ("InboundJob dead > 0", "즉시", "last_error 확인 후 원인 제거·재처리"),
            ("oldest pending job > 10분", "높음", "worker heartbeat/DB lock/HubSpot 확인"),
            ("send_failed 급증/SMTP auth", "즉시", "발송 중지, 자격증명/쿼터 확인"),
            ("post-sync 8회 소진", "높음", "HubSpot/Sheets 수동 동기화"),
            ("Sheet pending > 30분", "높음", "OAuth 권한·중복·API quota 확인"),
            ("poller 마지막 성공 > 20분", "높음", "worker task와 HubSpot search 확인"),
            ("worker heartbeat missing/stale", "즉시", "해당 background task·DB·프로세스 재시작 확인"),
            ("일일 발송량 80%", "주의", "캠페인/오류 트래픽 점검"),
        ],
        [3300, 1260, 4800],
        font_size=8.5,
    )
    add_heading(doc, "백업", 2)
    add_bullets(
        doc,
        [
            "PostgreSQL 일일 backup과 PITR, 보존 14~30일, 분기별 restore drill을 운영한다.",
            "migration 전 snapshot을 만든다.",
            "SESSION_SECRET/INTERNAL_API_TOKEN/OAuth client secret/GOOGLE_TOKEN_ENCRYPTION_KEY는 DB와 별도의 secret manager에 보관한다.",
            "Google grant 복구에는 원래 token encryption key가 필요하므로 키 회전 전 재연결 또는 복호화 migration을 계획한다.",
            "현재 /logs는 메모리 500건 ring buffer이므로 중앙 로그·에러 추적을 별도 연결한다.",
        ],
    )
    add_page_break(doc)
    add_heading(doc, "15. 장애 대응 Runbook", 1)
    add_image(
        doc,
        ASSET_DIR / "recovery_flow.png",
        "그림 2. 장애 복구 시 안전한 재처리 판단",
        "dead 문의, 확정 발송 실패, 발송 확인 필요, 외부 동기화 실패를 서로 다른 방식으로 복구하는 흐름",
        width=5.9,
    )
    add_table(
        doc,
        ["증상", "먼저 확인", "복구"],
        [
            ("문의가 안 잡힘", "stage ID, webhook 401, ticket body/association", "webhook 수정 후 poller 또는 test ticket 재이동"),
            ("auto_ack 없음", "Message auto_ack 상태, SMTP, first inbound 여부", "transient queue 관찰; failed면 원인 수정"),
            ("draft_failed", "LLM/정책 선택/HubSpot fetch log", "원인 수정 후 장애 복구에서 job 재처리"),
            ("승인 후 안 감", "approved/sending/send_failed, SMTP quota", "worker와 자격증명 수정; 반복 클릭 금지"),
            ("delivery_unknown", "Sent mailbox, provider log, Message-ID", "확인 후 수동 상태 결정; 바로 재전송 금지"),
            ("메일 sent, CRM 미반영", "post_send_sync_error/attempts", "외부 연결 복구, sync retry 또는 수동 이동"),
            ("Slack 반복/원치 않음", "SLACK_ENABLED, APPROVAL_CHANNEL, 중복 배포", "false/none로 재시작, old process 종료"),
            ("Sheets sync 실패", "OAuth account/editor, callback, schema headers", "권한 복구 후 영속 sync 재요청"),
            ("OAuth grant 복호화 실패", "GOOGLE_TOKEN_ENCRYPTION_KEY 변경 여부", "원래 key 복구 또는 재연결"),
        ],
        [2400, 3480, 3480],
        font_size=8.2,
    )

    add_page_break(doc)
    add_heading(doc, "16. 실서비스 전환 체크리스트", 1)
    add_table(
        doc,
        ["단계", "완료 조건"],
        [
            ("1. 데이터", "Managed PostgreSQL, migration, backup/PITR, restore test"),
            ("2. 보안", "Google OAuth UI, 최소 allowlist, 강한 secrets, webhook signature, private proxy"),
            ("3. HubSpot", "실제 stage ID 7종, webhook 2종, test ticket end-to-end"),
            ("4. 이메일", "override 주소 1건, 실제 From/SPF/DKIM/DMARC, rate/day cap"),
            ("5. Sheets", "전용 업무 Google 계정 OAuth, 세 탭 보존 확인, duplicate audit"),
            ("6. 정책", "auto_ack·서명·활성 지식 문서 승인, 개인정보 검토"),
            ("7. 운영", "delivery_unknown/dead/sync 실패 담당자와 SLA, 중앙 로그/경보"),
            ("8. 파일럿", "AUTO_SEND_THRESHOLD=1.01 유지, 소량 문의 사람 승인 운영"),
            ("9. 확대", "품질지표·오답·중복·실패율 검토 후 자동승인 여부 결정"),
        ],
        [1800, 7560],
    )
    add_callout(
        doc,
        "Go-live 게이트",
        "실제 고객에게 보내기 전 SEND_OVERRIDE_EMAIL이 비어 있는지, Slack이 의도한 값인지, 개발 stage 1→2가 실제 ID로 교체됐는지를 두 사람이 교차 확인한다.",
        fill=SUCCESS,
        accent=PERSO_TEAL,
    )

    add_page_break(doc)
    add_heading(doc, "17. 현재 부분 구현과 후속 백로그", 1)
    add_table(
        doc,
        ["요구", "현재 제공", "남은 구현"],
        [
            ("Gemini 회의록", "수동 interaction+URL", "Meet transcript ingestion/요약 connector"),
            ("카카오·전화", "수동 기록", "공식 API/CTI connector"),
            ("Stripe·PortOne", "계약·결제 URL/상태 저장", "결제 링크·Invoice 생성 및 webhook"),
            ("Reminder", "30일 무응답/14일 미접촉 대상 탐지", "승인 기반 reminder 발송"),
            ("리포트", "수동 /run/report와 이메일", "daily/weekly scheduler·Slack 제외 정책 유지"),
            ("수평 확장", "단일 web/worker", "dedicated workers, global rate limit, Sheets lock"),
            ("발송 exactly-once", "lease+delivery_unknown", "idempotency/delivery event 지원 메일 API"),
            ("대규모 Sheets", "5,000행 import, O(n) 조회", "pagination, batch, index/cache, duplicate audit"),
        ],
        [2400, 3600, 3360],
        font_size=8.4,
    )

    add_page_break(doc)
    add_heading(doc, "18. 상태와 주요 URL 부록", 1)
    add_heading(doc, "Message 상태", 2)
    add_table(
        doc,
        ["상태", "뜻", "운영 처리"],
        [
            ("drafting", "AI 초안 작성 중", "대기, 장기화 시 worker/log"),
            ("draft_failed", "초안 작성 실패", "원인 확인·재처리"),
            ("pending_approval", "사람 검토 필요", "수정·번역·미리보기·승인"),
            ("approved", "발송 큐 대기", "중복 클릭 금지"),
            ("sending:<worker>", "발송 lease 보유", "15분 전에는 개입하지 않음"),
            ("sent", "고객 채널 성공", "post-sync만 별도 확인"),
            ("test_sent", "내부 테스트 주소로만 발송", "CRM·단계·Sheets가 변하지 않았는지 확인"),
            ("send_failed", "확정 실패", "자격증명/수신자 확인"),
            ("delivery_unknown", "성공 여부 불명", "외부 로그 확인 전 재발송 금지"),
            ("rejected", "운영자가 발송 취소", "감사 이력 확인"),
        ],
        [2100, 3840, 3420],
        font_size=8.4,
    )
    doc.add_page_break()
    add_heading(doc, "주요 운영 URL", 2)
    add_table(
        doc,
        ["URL", "용도"],
        [
            ("/healthz", "DB readiness"),
            ("/internal/healthcheck", "토큰 보호 live integration 점검"),
            ("/webhooks/hubspot", "HubSpot 공개 webhook"),
            ("/messages", "답변 검토"),
            ("/customers", "고객·계약·통합 이력"),
            ("/operations", "문의/국가/후속 대상 인사이트"),
            ("/operations/recovery", "dead job·실패 발송·불명확 발송·외부 sync 복구"),
            ("/pipeline", "문의 단위 Kanban과 Sheets 연결"),
            ("/settings/users", "Google OAuth 사용자 승인"),
        ],
        [3000, 6360],
        font_size=8.0,
    )
    return save_document(doc, "PERSO_개발자_설정_아키텍처_가이드.docx")


def build_operator_guide() -> Path:
    doc = _configure_document(
        "PERSO 인바운드 자동화 - 비개발자 운영 가이드",
        "비개발자 기능·운영",
    )
    add_cover(
        doc,
        kicker="Operations Playbook",
        title="비개발자 기능 및\n운영 가이드",
        subtitle="문의 접수부터 답변, 고객 이력, 파이프라인, 시트까지",
        audience="대상: 세일즈·CS·Revenue Operations·팀 리더",
    )

    add_page_break(doc)
    add_heading(doc, "1. 5분 만에 이해하기", 1)
    add_para(
        doc,
        "이 시스템은 HubSpot에 들어온 신규 문의를 놓치지 않고, 고객에게 접수확인을 먼저 보낸 뒤 내부 정책을 참고한 답변 초안을 준비한다. 운영자는 초안을 검토해 실제 메일을 보내고, 고객 상태·HubSpot·Google Sheets를 한 화면에서 이어 관리한다.",
    )
    add_table(
        doc,
        ["자동으로 하는 일", "사람이 확인하는 일"],
        [
            ("HubSpot New 문의 수신과 누락 보완", "문의 본문과 고객 정보가 맞는지"),
            ("첫 문의 접수확인 메일", "상세 답변의 정확성·톤·민감정보"),
            ("AI 분류·정책 선택·한국어 초안", "외국어 번역과 최종 미리보기"),
            ("발송 후 HubSpot/시트 단계 갱신", "계약·결제·회의·카카오·전화 수동 기록"),
            ("답장 누락·무응답·갱신 임박 탐지", "어떤 고객에게 언제 후속 연락할지"),
        ],
        [4680, 4680],
    )
    add_callout(
        doc,
        "현재 안전 기본값",
        "상세 답변은 자동으로 바로 나가지 않는다. 사람이 '검토 완료 · 발송'을 눌러야 한다. Slack도 기본값은 꺼져 있다.",
        fill=SUCCESS,
        accent=PERSO_TEAL,
    )
    add_status_legend(doc)

    add_page_break(doc)
    add_heading(doc, "2. 메뉴 구조", 1)
    add_image(
        doc,
        ASSET_DIR / "menu_map.png",
        "그림 1. 왼쪽 메뉴의 세 가지 업무 영역",
        "인바운드 답장, 인사이트, 파이프라인 연동관리의 하위 메뉴 구조",
    )
    add_table(
        doc,
        ["메뉴", "언제 쓰나"],
        [
            ("인바운드 답장", "새 문의를 확인하고 답변을 검토하거나 정책·서명을 관리할 때"),
            ("인사이트", "문의량·국가 추이와 후속 연락이 필요한 고객을 찾을 때"),
            ("파이프라인 연동관리", "문의 단계를 옮기고 Google Sheets·견적·계약 흐름을 관리할 때"),
        ],
        [3000, 6360],
    )
    add_heading(doc, "내 계정의 역할", 2)
    add_table(
        doc,
        ["역할", "할 수 있는 일"],
        [
            ("조회 전용", "화면과 데이터를 읽고 검색·필터만 사용"),
            ("운영자", "답변 발송, 고객·계약·파이프라인 수정, 장애 복구"),
            ("관리자", "운영자 기능 + 정책·서명·Google 연결·사용자·로그 관리"),
        ],
        [2200, 7160],
    )

    add_page_break(doc)
    add_heading(doc, "3. 문의가 들어와 답변이 나가기까지", 1)
    add_image(
        doc,
        ASSET_DIR / "nondev_flow.png",
        "그림 2. 신규 문의 자동 처리 흐름",
        "HubSpot New 문의, 즉시 접수확인, AI와 내부 규정, 사람 검토, 실제 발송, 상태 동기화의 6단계 흐름",
    )
    add_heading(doc, "접수확인과 상세답변의 차이", 2)
    add_table(
        doc,
        ["구분", "접수확인", "상세답변"],
        [
            ("목적", "문의가 도착했고 곧 답하겠다고 알림", "정책과 상황에 맞는 실제 답변"),
            ("승인", "첫 문의에 자동", "기본적으로 사람 검토"),
            ("Slack", "보내지 않음", "초안이 검토 대기일 때만 가능"),
            ("단계 이동", "하지 않음", "성공 후 Meeting link sent"),
        ],
        [1560, 3900, 3900],
        font_size=8.6,
    )

    add_page_break(doc)
    add_heading(doc, "4. 매일 운영 순서", 1)
    add_steps(
        doc,
        [
            "문의 대시보드에서 오늘 접수, 작성 중, 검토 대기, 발송 실패를 확인한다.",
            "답변 검토에서 승인 대기 초안을 열고 고객 원문·번역·이력·정책 근거를 확인한다.",
            "제목·본문을 고치고 서명을 선택한다. 외국어면 번역 후 미리보기를 확인한다.",
            "검토 완료 · 발송을 누른 뒤 상태가 발송됨이 되는지 확인한다.",
            "파이프라인에서 해당 문의가 Meeting link sent로 이동했는지 본다.",
            "인사이트의 답장 누락·30일 무응답·14일 미접촉·60일 이내 갱신을 확인한다.",
            "장애 복구에서 dead 작업, 발송 실패, 발송 확인 필요, 외부 동기화 실패가 없는지 확인한다.",
            "미팅·전화·카카오·계약처럼 자동으로 들어오지 않는 기록을 고객 상세에 추가한다.",
        ],
    )
    add_callout(
        doc,
        "반복 클릭 금지",
        "발송 승인 후 '발송 대기' 또는 '승인됨'이면 작업이 큐에 들어간 상태다. 버튼을 다시 누르지 말고 상태가 바뀌기를 기다린다.",
        fill=CAUTION,
        accent="9A6A00",
    )
    add_heading(doc, "하루 마감", 2)
    add_bullets(
        doc,
        [
            "승인 대기와 발송 실패가 남아 있지 않은지 확인한다.",
            "발송 확인 필요가 있으면 보낸메일함과 제공자 로그를 먼저 확인한다.",
            "다음 액션 날짜가 지난 고객과 갱신 임박 고객의 담당자를 정한다.",
            "계약이 성사됐다면 계약 상태·금액·날짜·URL을 저장해 수주 DB 동기화를 유도한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "5. 문의 대시보드", 1)
    add_image(
        doc,
        ASSET_DIR / "01_dashboard.png",
        "그림 3. 문의 대시보드",
        "검토 대기, 오늘 접수, 작성 중, 오늘 발송, 확인 필요와 최근 답변",
    )
    add_table(
        doc,
        ["카드", "뜻", "행동"],
        [
            ("검토 대기", "pending_approval 상세 초안", "가장 먼저 검토"),
            ("오늘 접수", "오늘 들어온 고객 문의", "급증·누락 여부 확인"),
            ("작성 중", "AI가 분석·초안을 준비 중", "장기화 시 관리자 로그"),
            ("오늘 발송", "접수확인을 제외한 상세 답변", "일일 한도와 비교"),
            ("확인 필요", "작성 실패·발송 실패·불명확 발송", "상세 상태에 맞는 runbook"),
        ],
        [1800, 3600, 3960],
        font_size=8.8,
    )
    add_para(doc, "대시보드는 약 30초마다 갱신된다. 최근 답변 표에는 자동 접수확인을 제외한 실제 상세 답변만 보인다.", size=9.5, color=MUTED)

    add_page_break(doc)
    add_heading(doc, "6. 답변 검토 - 실제 발송 전", 1)
    add_image(
        doc,
        ASSET_DIR / "02_messages.png",
        "그림 4. 답변 검토 목록",
        "상태와 채널 필터를 제공하는 답변 검토 목록 화면",
        width=5.4,
    )
    add_heading(doc, "초안 한 건을 검토하는 순서", 2)
    add_steps(
        doc,
        [
            "고객 원문과 한국어 번역을 함께 읽는다.",
            "같은 티켓의 문의·자동 접수확인·과거 회신을 확인한다.",
            "AI 대화 요약, 고객 요청사항, 처리경과, 같은 사람/회사 이력을 본다.",
            "제목과 한국어 본문을 직접 고친다.",
            "기본 텍스트, 활성 HTML 브랜드 서명, 서명 없음 중 하나를 고른다.",
            "외국어 문의는 번역하기를 눌러 실제 고객 언어로 바꾼다.",
            "미리보기에서 본문·링크·서명·이미지를 확인한다.",
            "검토 완료 · 발송을 누르고 확인창에서 최종 승인한다.",
        ],
        size=9.8,
    )
    add_callout(
        doc,
        "첫 상세 회신의 가격",
        "현재 운영 규칙은 첫 상세 회신에서 구체 금액을 보내지 않는 것이다. AI 초안과 최종 발송 경로 모두 가격 문장을 제한하므로 가격은 이후 회신·미팅·견적 과정에서 안내한다.",
        fill=CAUTION,
        accent="9A6A00",
    )
    add_page_break(doc)
    add_heading(doc, "7. 답변 상태 읽는 법", 1)
    add_table(
        doc,
        ["화면 표시", "시스템 상태", "운영자가 할 일"],
        [
            ("AI 작성중", "drafting", "보통은 대기. 오래 지속되면 관리자 로그 확인"),
            ("작성 실패", "draft_failed", "AI/HubSpot/정책 오류 확인, 관리자 재처리"),
            ("승인 대기", "pending_approval", "수정·번역·미리보기 후 발송 또는 거절"),
            ("발송 대기", "approved", "중복 클릭 금지, worker 대기"),
            ("발송 중", "sending:<worker>", "잠시 대기"),
            ("발송됨", "sent", "외부 단계가 늦어도 메일 재발송 금지"),
            ("테스트 발송됨", "test_sent", "내부 주소로만 전송. 고객·CRM·시트는 변경되지 않음"),
            ("발송 실패", "send_failed", "수신자·SMTP·쿼터 확인"),
            ("발송 확인 필요", "delivery_unknown", "보낸메일함/제공자 로그 확인 전 재발송 금지"),
            ("거절", "rejected", "발송하지 않기로 한 기록"),
            ("수신", "received", "고객에게서 들어온 메시지"),
        ],
        [1900, 1900, 5560],
        font_size=8.6,
    )
    add_callout(
        doc,
        "발송 확인 필요는 실패가 아니다",
        "서버가 발송 도중 중단되어 SMTP가 받았는지 확정할 수 없다는 뜻이다. 같은 메일을 바로 다시 보내면 고객이 두 번 받을 수 있다. 관리자에게 확인을 요청한다.",
        fill=RISK,
        accent="9B1C1C",
    )

    add_page_break(doc)
    add_heading(doc, "8. 고객·회사 히스토리", 1)
    add_heading(doc, "고객 상세에서 관리하는 것", 2)
    add_table(
        doc,
        ["영역", "입력·확인 항목"],
        [
            ("기본 상태", "Negotiation/서비스 이용/기존 Pool/Lost, pipeline, 리드 온도"),
            ("고객 정보", "회사, 성함, 이메일, 전화번호, 국가, 산업군, user-seq, MQL/PQL"),
            ("다음 행동", "다음 액션, 예정일, 유입 소스, 운영 메모, Lost 사유"),
            ("통합 이력", "HubSpot 이메일·Deal·메모, 미팅·카카오·전화·수동 기록"),
            ("자료", "Gemini 회의록·Invoice·계약서 등 외부 URL"),
            ("계약", "금액, 날짜, 만료, 언어쌍, 단가, Invoice/결제 URL"),
        ],
        [2100, 7260],
    )
    add_steps(
        doc,
        [
            "고객 목록에서 회사·이름·이메일을 검색하거나 상태로 좁힌다.",
            "고객 상세의 HubSpot 동기화를 눌러 연락처, 최근 이메일 20건, Deal, 최신 메모를 가져온다.",
            "온라인 미팅·카카오·전화·계약서처럼 자동 연결이 없는 기록은 기록 추가로 남긴다.",
            "회의록·Invoice·계약서 URL을 artifact URL에 붙인다.",
            "미팅 기록은 Negotiation으로 이동할 수 있고, 계약 상태 active는 서비스 이용중으로 이동한다.",
            "계약 저장 때 실제 성사된 문의를 선택한다. 같은 고객의 다른 문의 카드는 그대로 유지된다.",
        ],
    )
    add_callout(
        doc,
        "같은 사람·같은 회사",
        "같은 이메일은 한 사람의 이력으로, 회사 도메인이 같으면 회사 이력으로 묶는다. Gmail·Naver 같은 개인 메일 도메인은 서로 다른 고객이 섞이지 않도록 회사로 묶지 않는다.",
    )
    add_para(doc, "현재 Gemini 회의록, 카카오, 전화, Stripe 데이터를 자동으로 가져오는 연결은 없다. 자료 URL과 요약을 사람이 기록한다.", color=MUTED, size=9.7)

    add_page_break(doc)
    add_heading(doc, "9. 이메일 규칙·서명과 정책·지식 문서", 1)
    add_image(
        doc,
        ASSET_DIR / "07_email_templates.png",
        "그림 5. 이메일 규칙·서명 목록",
        "자동 접수확인과 텍스트·HTML 서명을 관리하는 이메일 템플릿 목록",
        width=5.75,
    )
    add_table(
        doc,
        ["키", "용도", "주의"],
        [
            ("auto_ack", "첫 문의 즉시 접수확인", "이름 자리는 정확히 {name}"),
            ("signature_ko", "한국어 초안의 기본 텍스트 서명", "활성 상태만 적용"),
            ("signature_html_*", "검토 화면의 브랜드 서명", "공개 HTTPS 이미지 URL"),
            ("signature_html_hyeram", "기본 브랜드 HTML 서명", "키 이름 변경 금지"),
        ],
        [2200, 3600, 3560],
        font_size=8.7,
    )
    add_bullets(
        doc,
        [
            "기존 주요 키는 코드가 찾는 이름이므로 임의로 삭제하거나 새 이름으로 바꾸지 않는다.",
            "HTML 서명은 미리보기에서 확인하고, 이미지에 로컬 파일 경로나 로그인이 필요한 URL을 쓰지 않는다.",
            "여러 사람이 동시에 이메일 템플릿을 수정하면 마지막 저장이 이길 수 있으므로 한 명씩 작업한다.",
            "화면의 임의 {{변수}} 표시는 범용 치환 기능이 아니다. 새 키를 만든 것만으로 자동 발송에 연결되지 않는다.",
        ],
        size=10.1,
    )

    add_page_break(doc)
    add_heading(doc, "10. 정책·지식 문서 운영", 1)
    add_image(
        doc,
        ASSET_DIR / "06_knowledge.png",
        "그림 6. AI가 참고하는 정책·지식 문서 목록",
        "활성·비활성, 답변 분류, 적용 범위, 버전을 보여주는 정책 지식 문서 화면",
    )
    add_steps(
        doc,
        [
            "새 문서에서 제목과 변경 불가한 영문 문서 키를 정한다.",
            "적용 범위를 inbound 또는 both로 둔다.",
            "답변 분류, 검색 태그, AI 선택 요약을 명확히 쓴다.",
            "정책 본문에는 확정 기준, 예외, 고객에게 써도 되는 표현을 적는다.",
            "검토가 끝난 문서만 활성으로 저장하고 변경 사유를 남긴다.",
        ],
    )
    add_bullets(
        doc,
        [
            "활성 문서만 AI가 사용한다. spam 문의에는 정책 문서를 제공하지 않는다.",
            "수정할 때 다른 사용자가 먼저 저장하면 충돌 안내가 뜬다. 최신 내용을 다시 확인한다.",
            "변경 이력에는 수정 전 전체 내용, 편집자, 사유가 남는다.",
            "삭제 기능 대신 비활성 상태를 사용한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "11. 인사이트와 업데이트 필요 고객", 1)
    add_image(
        doc,
        ASSET_DIR / "03_operations.png",
        "그림 7. 월별 문의량·누적 문의와 국가 추이",
        "일 월 년 전환, 선택 기간 문의, 누적 문의, 평균 품질, 문의량과 누적 추이, 국가별 문의 추이 화면",
        width=4.8,
    )
    add_table(
        doc,
        ["선택", "보는 기간", "차트 의미"],
        [
            ("일", "최근 30일", "막대=그날 문의, 선=누적 문의"),
            ("월", "최근 12개월", "막대=그달 문의, 선=누적 문의"),
            ("년", "최근 5년", "막대=그해 문의, 선=누적 문의"),
        ],
        [1560, 3000, 4800],
    )
    add_table(
        doc,
        ["업데이트 대상", "기준", "자동 발송 여부"],
        [
            ("답장 누락 가능성", "최근 고객 문의가 우리 답변보다 새롭거나 답변 없음", "아니오"),
            ("30일 고객 답장 없음", "우리 마지막 답변 후 고객 회신 없음", "아니오"),
            ("14일 이상 소통 없음", "Negotiation 최근 활동이 14일 이전", "아니오"),
            ("60일 이내 계약 만료", "active 계약의 만료일이 60일 안", "아니오"),
            ("Closed Lost", "종료 사유 확인 대상", "아니오"),
            ("업셀 후보", "서비스 이용 중, Business/Enterprise 아님", "아니오"),
        ],
        [2400, 4560, 2400],
        font_size=8.3,
    )
    add_para(
        doc,
        "알림판 원칙: 후속 대상을 찾아줄 뿐 Reminder 메일은 자동 발송하지 않는다. 담당자가 맥락을 보고 다음 액션을 정한다.",
        size=9.2,
        color=MUTED,
        after=0,
    )

    add_page_break(doc)
    add_heading(doc, "12. 문의 파이프라인", 1)
    add_image(
        doc,
        ASSET_DIR / "04_pipeline.png",
        "그림 8. 문의 단위 Kanban 파이프라인",
        "New, Meeting link sent, Negotiation, Contracted, Onboarding, Active, Closed Lost 열로 구성된 문의 파이프라인",
        width=5.1,
    )
    add_image(
        doc,
        ASSET_DIR / "status_flow.png",
        "그림 9. 문의 단계 흐름",
        "New에서 Meeting link sent, Negotiation, Contracted, Onboarding, Active로 이어지고 Closed Lost로 종료되는 단계",
        width=5.1,
    )
    add_bullets(
        doc,
        [
            "카드 한 장은 고객 한 명이 아니라 문의/대화 한 건이다. 같은 고객의 문의가 여러 건이면 카드도 여러 장이다.",
            "카드를 열 사이로 끌거나 카드의 단계 선택 상자를 바꾼다.",
            "먼저 사이트가 저장되고, Ticket ID와 Client ID가 있으면 HubSpot과 Sheets도 갱신한다.",
            "열 위에 HubSpot ID가 보이면 실제 단계 연결, 로컬 전용이면 사이트 안에서만 이동한다.",
            "일부 외부 시스템 동기화 실패가 보이면 사이트 저장은 성공한 것이므로 고객에게 메일을 다시 보내지 않는다.",
        ],
        size=9.8,
    )

    add_page_break(doc)
    add_heading(doc, "13. Google Sheets 연결과 동기화", 1)
    add_image(
        doc,
        ASSET_DIR / "05_integrations.png",
        "그림 10. 파이프라인 아래 Google Sheets 연결 패널",
        "Inbound DB, Inbound 퀄리티 분석, 수주 DB 보존 원칙과 callback URI, 연결·동기화 버튼이 있는 Google Sheets 패널",
        width=5.3,
    )
    add_heading(doc, "관리자가 처음 한 번", 2)
    add_steps(
        doc,
        [
            "화면에 OAuth 설정 필요가 보이면 개발자가 Google OAuth Client ID/Secret을 서버에 설정한다.",
            "내 Google 계정 연결을 눌러 대상 시트를 편집할 수 있는 업무 계정으로 동의한다.",
            "연결된 계정 이메일이 화면에 보이는지 확인한다.",
            "시트 읽기·동기화를 누르면 요청이 먼저 접수된다. 처리 중 표시가 완료 또는 실패로 바뀌는지 확인한다.",
        ],
        size=9.8,
    )
    add_table(
        doc,
        ["동작", "자동/수동", "시트 영향"],
        [
            ("신규 문의", "자동+재시도", "Inbound DB Client ID 행 생성/갱신"),
            ("상세 답변·단계 이동", "자동", "해당 Inbound 행의 Deal Stage만 갱신"),
            ("contracted/active 계약", "자동+재시도", "수주 DB 행 생성/갱신"),
            ("기존 Inbound DB 가져오기", "영속 동기화 요청", "화면을 나가도 background에서 읽고 상태 표시"),
            ("Inbound 퀄리티 분석", "앱에서 쓰지 않음", "수식·차트 그대로 보존"),
            ("형식 보호", "항상", "매핑된 DB 셀만 쓰고 헤더·수식·서식은 보존"),
        ],
        [2600, 2600, 4160],
        font_size=7.9,
    )

    add_heading(doc, "14. 계약·결제와 견적", 1)
    add_image(
        doc,
        ASSET_DIR / "08_quote.png",
        "그림 11. 한국어 Business Plan 견적 계산기",
        "월 사용량, 언어 수, 더빙 종류로 추천 Tier와 최종 크레딧을 계산하는 화면",
        width=5.0,
    )
    add_heading(doc, "계약 기록", 2)
    add_table(
        doc,
        ["분류", "입력 항목"],
        [
            ("기본", "플랜, 상태, 금액, 통화, 계약일, 만료일"),
            ("연결 문의", "실제로 계약으로 성사된 문의 선택, 해당 Client ID 스냅샷"),
            ("결제", "포트원/Stripe/계좌이체, 일시불/할부, 예정일, 실제 입금일"),
            ("상품", "언어쌍, 단가, 크레딧, 초대/Queue/동시작업/Space 한도"),
            ("문서", "견적서, Invoice, 결제 링크, 계약서 방식"),
            ("수주 DB", "Billing Email, owner, space_seq, plan start, Enterprise Name 등"),
        ],
        [1800, 7560],
        font_size=8.4,
    )
    add_bullets(
        doc,
        [
            "contracted 저장 시 선택한 문의만 Contracted로 이동하고 그 문의의 Client ID로 수주 DB 동기화를 시도한다.",
            "계약 금액은 소수점까지 정확하게 저장한다. 천 단위 쉼표를 입력해도 된다.",
            "active 저장 시 서비스 이용중·Active로 이동하고 수주 DB를 동기화한다.",
            "견적 계산기 결과는 메일·계약·Invoice에 자동 저장되지 않는다. 최종 조건은 계약 화면에 따로 기록한다.",
            "이 시스템은 Stripe/포트원 결제 링크나 Invoice를 실제로 생성하지 않는다. 만들어진 URL과 결제 사실을 관리한다.",
            "입금 확인은 자동 연결되지 않았으므로 실제 입금일과 상태를 사람이 갱신한다.",
        ],
        size=9.5,
    )

    add_heading(doc, "15. Slack", 1)
    add_table(
        doc,
        ["질문", "답"],
        [
            ("Slack은 언제 오나?", "상세 AI 초안이 승인 대기가 된 순간에만 가능"),
            ("접수확인 때문에 오나?", "아니오"),
            ("같은 초안이 반복되나?", "성공하면 1회 기록. 실패한 알림만 제한 재시도"),
            ("Slack을 즉시 끄려면?", "관리자가 SLACK_ENABLED=false, APPROVAL_CHANNEL=none으로 재시작"),
        ],
        [3300, 6060],
        font_size=8.9,
    )
    add_callout(
        doc,
        "개인정보",
        "Slack 승인 카드에는 고객 이메일, 문의 일부, 초안이 포함될 수 있다. 비공개 승인 채널을 사용하고 최소 인원만 접근하도록 한다.",
        fill=CAUTION,
        accent="9A6A00",
    )

    add_page_break(doc)
    add_heading(doc, "16. 문제 상황별 해결", 1)
    add_image(
        doc,
        ASSET_DIR / "recovery_flow.png",
        "그림 12. 장애 복구 화면에서 선택하는 안전한 처리",
        "문의 작업·발송 실패·발송 확인 필요·외부 동기화 실패별 복구 원칙",
        width=5.9,
    )
    add_table(
        doc,
        ["상황", "내가 먼저 할 일", "개발자/관리자에게 전달"],
        [
            ("문의가 안 보임", "HubSpot이 New인지, 본문·연락처가 있는지", "티켓 번호와 발생 시각"),
            ("접수확인은 갔는데 초안 없음", "답변 검토의 작성 실패·장애 복구 확인", "화면 상태와 로그 시각"),
            ("승인했는데 발송 대기", "반복 클릭하지 않고 잠시 대기", "5분 이상 지속한 메시지 번호"),
            ("발송 실패", "수신자 주소 확인 후 장애 복구에서 재시도", "메시지 번호, 오류 시각"),
            ("발송 확인 필요", "보낸메일함 확인 후 장애 복구에서 결과 확정", "Message-ID 확인 결과"),
            ("메일은 갔는데 단계가 안 변함", "파이프라인과 연결 패널 확인", "HubSpot/Sheets 중 실패한 대상"),
            ("Slack이 자꾸 옴", "같은 초안인지, 이전 배포가 켜져 있는지", "알림 시각·초안 번호"),
            ("Sheets 동기화 실패", "연결 계정 확인 후 다시 요청", "장애 복구·화면 오류와 callback URI"),
            ("계약이 수주 DB에 없음", "contracted/active와 Client ID 확인", "고객·계약 번호"),
        ],
        [2280, 3960, 3120],
        font_size=8.2,
    )
    add_callout(
        doc,
        "운영 로그",
        "관리자 /logs는 최근 문제를 보는 보조 화면이고 서버 재시작 시 사라진다. 화면을 비우기 전에 필요한 시각과 메시지 번호를 기록한다.",
    )

    add_page_break(doc)
    add_heading(doc, "17. 실제 계정 연결 전후 역할", 1)
    add_table(
        doc,
        ["작업", "비개발자 관리자", "개발자"],
        [
            ("정책·서명 검토", "최종 문구와 활성 상태 승인", "seed/키 연결 확인"),
            ("HubSpot", "실제 단계 이름·운영 규칙 결정", "token, webhook, 내부 stage ID 설정"),
            ("Google Sheets", "편집 가능한 전용 계정으로 동의", "OAuth client/callback/secret 설정"),
            ("SMTP", "From 이름·주소·테스트 수신 확인", "자격증명, SPF/DKIM/DMARC, quota"),
            ("Slack", "비공개 채널·수신자 결정", "flag/token/channel 설정"),
            ("권한", "승인할 사용자·역할 결정", "Google OAuth/allowlist/session 설정"),
            ("운영 시작", "소량 문의 전건 사람 승인", "모니터링·백업·장애 대응"),
        ],
        [2200, 3580, 3580],
        font_size=8.5,
    )
    add_heading(doc, "현재 계정 연결 상황", 2)
    add_bullets(
        doc,
        [
            "HubSpot은 개발자 계정의 New(1)→Waiting on contact(2)를 임시 사용한다.",
            "Google Sheets는 서비스 계정 공유가 막혀 있고 사용자 OAuth Client ID/Secret 설정이 남아 있다.",
            "Slack은 꺼져 있어 초안이 준비돼도 보내지 않는다.",
            "상세 답변 자동승인은 꺼져 있고 사람이 승인한다.",
        ],
    )

    add_page_break(doc)
    add_heading(doc, "18. 실무용 한 장 체크리스트", 1)
    add_heading(doc, "답변을 보내기 전", 2)
    add_bullets(
        doc,
        [
            "□ 고객 원문과 번역을 함께 읽었다.",
            "□ 같은 사람·회사·티켓의 이전 대화를 확인했다.",
            "□ 내부 정책과 실제 제품 범위를 벗어난 약속이 없다.",
            "□ 첫 상세 회신에 구체 가격을 넣지 않았다.",
            "□ 외국어 답변은 번역 후 미리보기로 확인했다.",
            "□ 링크와 서명 이미지가 공개 HTTPS로 열린다.",
            "□ 수신자와 제목이 맞다.",
        ],
    )
    add_heading(doc, "발송 후", 2)
    add_bullets(
        doc,
        [
            "□ 메시지 상태가 발송됨이다.",
            "□ 문의 단계가 Meeting link sent로 이동했다.",
            "□ HubSpot/Sheets 일부 실패가 있으면 메일을 재발송하지 않고 외부 동기화만 확인했다.",
            "□ 다음 액션과 날짜를 고객 상세에 기록했다.",
            "□ 미팅·전화·카카오·계약 자료를 통합 이력에 남겼다.",
        ],
    )
    add_heading(doc, "절대 하지 않을 것", 2)
    add_bullets(
        doc,
        [
            "□ 발송 대기 중 버튼 반복 클릭",
            "□ 발송 확인 필요를 확인 없이 재발송",
            "□ Google 시트 헤더·수식·분석 탭 구조 임의 변경",
            "□ 비밀키·고객 개인정보를 공개 채널이나 문서에 복사",
        ],
    )
    add_para(
        doc,
        "한 문장 요약: 반복 작업과 누락 탐지는 자동화하고, 고객에게 나가는 최종 판단은 사람이 책임진다.",
        size=9.2,
        bold=True,
        color=PERSO_TEAL,
        before=2,
        after=0,
    )

    add_page_break(doc)
    add_heading(doc, "19. 용어 사전", 1)
    add_table(
        doc,
        ["용어", "쉬운 뜻"],
        [
            ("Webhook", "HubSpot에서 일이 생기자마자 사이트에 알려주는 신호"),
            ("Polling", "Webhook을 놓쳤을 때 사이트가 10분마다 다시 확인하는 방법"),
            ("auto_ack", "문의 접수 사실만 먼저 알리는 자동 메일"),
            ("AI 초안", "고객에게 바로 가지 않고 사람이 검토하는 한국어 답변 초안"),
            ("SMTP", "실제로 고객 메일함에 이메일을 보내는 서버"),
            ("Meeting link sent", "이 시스템에서는 상세 답변 발송이 끝난 단계"),
            ("Client ID", "사이트와 Google 시트의 같은 문의를 연결하는 번호"),
            ("MQL/PQL", "마케팅 또는 제품 사용 신호를 기준으로 한 리드 구분"),
            ("delivery_unknown", "메일 서버가 받았는지 확정할 수 없어 수동 확인이 필요한 상태"),
            ("동기화", "사이트의 정보를 HubSpot·시트와 맞추는 작업"),
        ],
        [2600, 6760],
        font_size=8.8,
    )
    add_para(
        doc,
        "문서 끝 - 화면이나 운영 규칙이 바뀌면 팀 리더와 개발자가 함께 이 가이드를 갱신한다.",
        size=9.5,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=12,
    )
    return save_document(doc, "PERSO_비개발자_기능_운영_가이드.docx")


def main() -> None:
    build_diagrams()
    paths = [build_developer_guide(), build_operator_guide()]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
